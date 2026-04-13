import io
import os
import re
import shutil
import subprocess
import tempfile
import uuid

from fastapi import HTTPException, UploadFile
from fastapi.responses import StreamingResponse

import schemas
from repositories.km_repository import KMRepository


class KMService:
    def __init__(self, repo: KMRepository):
        self.repo = repo

    def get_articles(self, skip: int = 0, limit: int = 100, category: str | None = None, search: str | None = None):
        return self.repo.get_all(skip=skip, limit=limit, category=category, search=search)

    def get_article(self, article_id: str):
        article = self.repo.get_by_id(article_id)
        if not article:
            raise HTTPException(status_code=404, detail="KM Article not found")
        return article

    def create_article(self, article_create: schemas.KMArticleCreate, author_id: int):
        return self.repo.create(article=article_create, author_id=author_id)

    def update_article(self, article_id: str, article_update: schemas.KMArticleUpdate):
        db_article = self.repo.get_by_id(article_id)
        if not db_article:
            raise HTTPException(status_code=404, detail="KM Article not found")

        update_data = article_update.model_dump(exclude_unset=True)
        try:
            return self.repo.update(db_article, update_data)
        except ValueError as e:
            raise HTTPException(status_code=409, detail=str(e))

    def delete_article(self, article_id: str):
        db_article = self.repo.get_by_id(article_id)
        if not db_article:
            raise HTTPException(status_code=404, detail="KM Article not found")
        self.repo.delete(db_article)
        return {"ok": True}

    def get_article_history(self, article_id: str):
        db_article = self.repo.get_by_id(article_id)
        if not db_article:
            raise HTTPException(status_code=404, detail="KM Article not found")
        return self.repo.get_history(article_id)

    # ──────────────────────────────────────────────
    # Word Export / Import
    # ──────────────────────────────────────────────

    def _chapter_depth(self, chapter_no: str | None) -> int:
        """Return heading level: '1' → 1, '1.1' → 2, '1.1.1' → 3 (capped at 3)."""
        if not chapter_no:
            return 2
        clean = chapter_no.strip().rstrip('.0')
        dots = clean.count('.')
        return min(dots + 2, 4)  # book title uses h1; chapters start at h2

    def _quill_html_to_docx_html(self, html_content: str, chapter_prefix: str) -> str:
        """
        Convert Quill editor HTML into clean HTML suitable for htmldocx.

        Quill stores lists as flat <ol>/<ul> with ql-indent-N classes for nesting.
        This pre-processor:
          - Replaces <ol><li> items with paragraphs bearing explicit chapter-prefix
            numbering (e.g. "3.1", "3.2", "3.1.1" for indent-1)
          - Replaces <ul><li> items with bullet paragraphs (○ for top-level, - for indent-1)
          - Marks paragraphs with data-indent-level for post-processing by python-docx
          - Strips Quill-specific CSS classes (ql-font-*, ql-size-*, etc.)
        """
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, 'html.parser')

        # Strip the trailing ".0" suffix from chapter prefix: "3.0" → "3", "10.0" → "10"
        # Use regex to only strip a final ".0" segment, NOT arbitrary trailing zeros
        raw = (chapter_prefix or '').strip()
        prefix = re.sub(r'\.0$', '', raw)

        # Counters for ordered list numbering
        ol_counter_0 = 0  # top-level
        ol_counter_1 = 0  # indent-1

        # Process all list containers (ol, ul) and convert to paragraphs
        for list_el in soup.find_all(['ol', 'ul']):
            is_ordered = list_el.name == 'ol'
            items = list_el.find_all('li', recursive=False)

            replacement_elements = []
            for li in items:
                # Determine indent level from ql-indent-N class
                indent = 0
                classes = li.get('class', [])
                for cls in classes:
                    if cls.startswith('ql-indent-'):
                        try:
                            indent = int(cls.replace('ql-indent-', ''))
                        except ValueError:
                            pass

                # Build the numbering prefix text and indent level marker
                marker = ''
                # indent_level: 1=top-ol, 2=indent-1-ol, 3=bullet-top, 4=bullet-indent-1
                indent_level = 1

                if is_ordered:
                    if indent == 0:
                        ol_counter_0 += 1
                        ol_counter_1 = 0  # reset sub-counter
                        if prefix:
                            marker = f"{prefix}.{ol_counter_0}\u00A0\u00A0"
                        else:
                            marker = f"{ol_counter_0}.\u00A0\u00A0"
                        indent_level = 1
                    elif indent == 1:
                        ol_counter_1 += 1
                        if prefix:
                            marker = f"{prefix}.{ol_counter_0}.{ol_counter_1}\u00A0\u00A0"
                        else:
                            marker = f"{ol_counter_0}.{ol_counter_1}\u00A0\u00A0"
                        indent_level = 2
                    else:
                        marker = f"({indent})\u00A0\u00A0"
                        indent_level = indent + 1
                else:
                    # Unordered list
                    if indent == 0:
                        marker = "○\u00A0"
                        indent_level = 3
                    elif indent == 1:
                        marker = "-\u00A0"
                        indent_level = 4
                    else:
                        marker = "•\u00A0"
                        indent_level = indent + 3

                # Get inner HTML of the <li> (preserving child tags like <span>, <strong>)
                inner_html = li.decode_contents()

                # Remove Quill-specific classes from inner content
                inner_html = re.sub(r'\bql-font-\w+', '', inner_html)
                inner_html = re.sub(r'\bql-size-\w+', '', inner_html)
                inner_html = re.sub(r'\bclass="[\s]*"', '', inner_html)

                # Carry over the li's inline style (line-height, etc.)
                li_style = li.get('style', '')

                # Use data-indent-level attribute for post-processing (htmldocx ignores margin-left)
                p_tag = soup.new_tag('p', style=li_style)
                p_tag['data-indent-level'] = str(indent_level)

                # Insert marker + TAB so Word tab stop controls text column
                marker_span = soup.new_tag('span')
                marker_span.string = marker.rstrip('\u00A0 ') + '\t'
                p_tag.append(marker_span)

                # Parse and append inner content
                inner_soup = BeautifulSoup(inner_html, 'html.parser')
                for child in list(inner_soup.children):
                    p_tag.append(child)

                replacement_elements.append(p_tag)

            # Replace the <ol>/<ul> with the generated <p> elements
            for p in replacement_elements:
                list_el.insert_before(p)
            list_el.decompose()

        # Clean up remaining Quill CSS classes in non-list content
        for tag in soup.find_all(True):
            classes = tag.get('class', [])
            if classes:
                cleaned = [c for c in classes if not c.startswith('ql-')]
                if cleaned:
                    tag['class'] = cleaned
                else:
                    del tag['class']

        return str(soup)

    # Indent layout matching kmArticle.css (em → cm at 12pt: 1em ≈ 0.423cm)
    #
    # CSS layout:
    #   --km-base-indent: 3em  (1.27cm) — ol text column
    #   --km-step: 1.5em       (0.63cm)
    #   ol top-level:  marker at 0,       text at 3em    (1.27cm)
    #   ul top-level:  ○ at 3em (1.27cm), text at 4.5em  (1.9cm)
    #   ol indent-1:   marker at 4.5em,   text at 8.5em  (3.6cm)
    #   ul indent-1:   - marker,          text at 8.5em  (3.6cm)
    #
    # Word implementation: hanging indent + tab stop at text column.
    # The marker text (e.g. "4.1\t") hangs left of left_indent; the tab
    # jumps to the tab stop = left_indent position; wrapped lines start
    # at left_indent. This guarantees text alignment regardless of marker
    # width.
    _INDENT_LAYOUT = {
        # key: (marker_start_cm, text_column_cm)
        'ol_0':  (0,    1.27),   # "4.1" at 0cm,     text at 1.27cm
        'ul_0':  (1.27, 1.9),    # "○"   at 1.27cm,  text at 1.9cm
        'ol_1':  (1.9,  3.6),    # "4.1.1" at 1.9cm, text at 3.6cm
        'ul_1':  (3.0,  3.6),    # "-"   at 3.0cm,   text at 3.6cm
    }

    def _apply_docx_indentation(self, doc, para_start_idx: int):
        """
        Post-process paragraphs added by htmldocx: detect list paragraphs by
        their marker prefix and apply Word hanging-indent + tab stop to match
        the web CSS layout from kmArticle.css.

        Steps per paragraph:
          1. Detect marker type (ol_0, ul_0, ol_1, ul_1)
          2. Set hanging indent (left_indent at text column, first_line pulls back)
          3. Add a tab stop at text column
          4. Insert a <w:tab/> element between marker run and content run
             (htmldocx drops \\t from HTML, so we must add the Word tab element)
        """
        from docx.shared import Cm
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Regex patterns for each marker type
        marker_patterns = {
            'ol_0': re.compile(r'^(\d+(?:\.\d+)\s*)'),       # "4.1 "
            'ol_1': re.compile(r'^(\d+(?:\.\d+){2,}\s*)'),    # "4.1.1 "
            'ul_0': re.compile(r'^(○\s*)'),                    # "○ "
            'ul_1': re.compile(r'^(-\s*)'),                    # "- "
            'ul_x': re.compile(r'^(•\s*)'),                    # "• "
        }

        for para in doc.paragraphs[para_start_idx:]:
            text = para.text.strip()
            if not text:
                continue

            layout_key = None
            marker_len = 0

            # Detect and measure marker — order matters (ol_1 before ol_0)
            m = marker_patterns['ol_1'].match(text)
            if m:
                layout_key, marker_len = 'ol_1', m.end()
            else:
                m = marker_patterns['ol_0'].match(text)
                if m:
                    layout_key, marker_len = 'ol_0', m.end()
                elif marker_patterns['ul_0'].match(text):
                    m = marker_patterns['ul_0'].match(text)
                    layout_key, marker_len = 'ul_0', m.end()
                elif marker_patterns['ul_1'].match(text):
                    m = marker_patterns['ul_1'].match(text)
                    layout_key, marker_len = 'ul_1', m.end()
                elif marker_patterns['ul_x'].match(text):
                    m = marker_patterns['ul_x'].match(text)
                    layout_key, marker_len = 'ul_1', m.end()

            if not layout_key or layout_key not in self._INDENT_LAYOUT:
                continue

            marker_cm, text_cm = self._INDENT_LAYOUT[layout_key]

            # 1. Hanging indent
            para.paragraph_format.left_indent = Cm(text_cm)
            para.paragraph_format.first_line_indent = Cm(marker_cm - text_cm)

            # 2. Tab stop at text column
            pPr = para._element.get_or_add_pPr()
            tabs = pPr.find(qn('w:tabs'))
            if tabs is None:
                tabs = OxmlElement('w:tabs')
                pPr.append(tabs)
            tab_el = OxmlElement('w:tab')
            tab_el.set(qn('w:val'), 'left')
            tab_el.set(qn('w:pos'), str(int(Cm(text_cm))))
            tabs.append(tab_el)

            # 3. Insert <w:tab/> between marker and content in the XML runs.
            #    htmldocx drops \t, so the marker run ends with a trailing
            #    space. We trim that space and insert a real <w:tab/> run.
            runs = para._element.findall(qn('w:r'))
            chars_seen = 0
            for r in runs:
                t_el = r.find(qn('w:t'))
                if t_el is None or not t_el.text:
                    chars_seen += len(r.text or '')
                    continue
                t_text = t_el.text
                # Check if this run contains the marker boundary
                run_start = chars_seen
                run_end = chars_seen + len(t_text)
                if run_start < marker_len <= run_end:
                    # Split: trim marker part, insert tab, keep rest
                    split_at = marker_len - run_start
                    marker_text = t_text[:split_at].rstrip()
                    rest_text = t_text[split_at:].lstrip()

                    # Update this run to just the marker text
                    t_el.text = marker_text

                    # Create a new run with <w:tab/> element
                    tab_run = OxmlElement('w:r')
                    # Copy run properties if they exist
                    rPr = r.find(qn('w:rPr'))
                    if rPr is not None:
                        tab_run.append(rPr.__class__(rPr))
                    tab_char = OxmlElement('w:tab')
                    tab_run.append(tab_char)

                    # Insert tab run right after the marker run
                    r.addnext(tab_run)

                    # If there's remaining text, create another run for it
                    if rest_text:
                        rest_run = OxmlElement('w:r')
                        if rPr is not None:
                            rest_run.append(rPr.__class__(rPr))
                        rest_t = OxmlElement('w:t')
                        rest_t.text = rest_text
                        rest_t.set(qn('xml:space'), 'preserve')
                        rest_run.append(rest_t)
                        tab_run.addnext(rest_run)
                    break
                chars_seen = run_end

    def export_docx(self, article_id: str) -> StreamingResponse:
        """Generate a .docx of the KM book and all its chapters using htmldocx."""
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from htmldocx import HtmlToDocx
        from urllib.parse import quote

        book = self.repo.get_by_id(article_id)
        if not book:
            raise HTTPException(status_code=404, detail="KM Article not found")

        # Resolve to book root
        root_id = book.parent_id or book.id
        root = self.repo.get_by_id(root_id)
        if not root:
            raise HTTPException(status_code=404, detail="KM root article not found")

        children = self.repo.get_children(root_id)

        # Sort chapters by chapter_no
        def sort_key(ch):
            parts = []
            for seg in (ch.chapter_no or '0').split('.'):
                try:
                    parts.append(int(seg))
                except ValueError:
                    parts.append(0)
            return parts

        chapters = sorted(children, key=sort_key) if children else [root]

        # Create Word document
        doc = DocxDocument()

        # Set default font: Calibri + 微軟正黑體 for CJK on all styles
        def _set_font(style_obj, size=None):
            """Set Calibri (Latin) + 微軟正黑體 (East Asian) on a style."""
            style_obj.font.name = 'Calibri'
            if size:
                style_obj.font.size = size
            rpr = style_obj.element.get_or_add_rPr()
            rFonts = rpr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = rpr.makeelement(qn('w:rFonts'), {})
                rpr.insert(0, rFonts)
            rFonts.set(qn('w:eastAsia'), '微軟正黑體')

        _set_font(doc.styles['Normal'], Pt(12))
        for hlevel in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
            if hlevel in doc.styles:
                _set_font(doc.styles[hlevel])

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2)

        # Title
        title_para = doc.add_heading(root.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # HTML-to-docx converter
        parser = HtmlToDocx()

        # Add chapters
        for ch in chapters:
            level = self._chapter_depth(ch.chapter_no)
            chapter_label = f"{ch.chapter_no + ' ' if ch.chapter_no else ''}{ch.title}"
            doc.add_heading(chapter_label, level=min(level, 4))

            content = ch.content or ''
            if content.strip():
                # Pre-process Quill HTML: convert <ol>/<ul> to numbered/bulleted
                # paragraphs with chapter-prefix numbering
                processed = self._quill_html_to_docx_html(content, ch.chapter_no or '')

                # Track paragraph count before adding content
                para_count_before = len(doc.paragraphs)

                # Wrap in proper HTML structure for htmldocx
                html_body = f"""<html><head><meta charset="utf-8"/></head>
                <body style="font-family: Calibri, '微軟正黑體', sans-serif; font-size: 12pt;">
                {processed}
                </body></html>"""
                try:
                    parser.add_html_to_document(html_body, doc)
                    # Post-process: apply Word-native indentation to list paragraphs
                    self._apply_docx_indentation(doc, para_count_before)
                except Exception:
                    # Fallback: add as plain text if HTML parsing fails
                    from bs4 import BeautifulSoup
                    text = BeautifulSoup(content, 'html.parser').get_text(separator='\n').strip()
                    if text:
                        for line in text.split('\n'):
                            line = line.strip()
                            if line:
                                doc.add_paragraph(line)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # ASCII-only fallback for filename (latin-1 safe)
        ascii_filename = re.sub(r'[^a-zA-Z0-9_\-]', '_', root.title)[:50] or 'km_export'
        # UTF-8 encoded filename for browsers that support RFC 5987
        encoded_name = quote(f"{root.title[:50]}.docx")
        return StreamingResponse(
            buffer,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f"attachment; filename=\"{ascii_filename}.docx\"; filename*=UTF-8''{encoded_name}"
            }
        )

    def import_docx(self, article_id: str, file: UploadFile, author_id: int) -> dict:
        """
        Parse an uploaded .docx and update matching chapters by chapter_no.
        Matching strategy: heading text starts with the chapter_no prefix.
        Returns a summary of updated chapters.
        """
        book = self.repo.get_by_id(article_id)
        if not book:
            raise HTTPException(status_code=404, detail="KM Article not found")

        root_id = book.parent_id or book.id
        root = self.repo.get_by_id(root_id)
        children = self.repo.get_children(root_id)
        all_chapters = children if children else [root]

        # Build a map: chapter_no → db article
        chapter_map = {
            (ch.chapter_no or '').strip(): ch
            for ch in all_chapters
            if ch.chapter_no
        }
        # Also include root (no chapter_no) for single-chapter books
        if not children:
            chapter_map['__root__'] = root

        # Save upload to temp file
        contents = file.file.read()
        if not contents:
            raise HTTPException(status_code=400, detail="Uploaded file is empty")

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, 'upload.docx')
            html_path = os.path.join(tmpdir, 'output.html')

            with open(docx_path, 'wb') as f:
                f.write(contents)

            # Convert docx → HTML with pandoc
            result = subprocess.run(
                ['pandoc', docx_path, '-o', html_path,
                 '--from=docx', '--to=html', '--wrap=none'],
                capture_output=True, text=True
            )
            if result.returncode != 0:
                raise HTTPException(status_code=400, detail=f"Could not parse Word file: {result.stderr}")

            with open(html_path, 'r', encoding='utf-8') as f:
                html = f.read()

        # Parse HTML and split by headings
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')

        # Walk top-level elements, split at h2/h3/h4 headings
        sections: list[dict] = []
        current: dict | None = None

        for el in soup.body.children if soup.body else []:
            tag = getattr(el, 'name', None)
            if tag in ('h2', 'h3', 'h4'):
                if current is not None:
                    sections.append(current)
                current = {'heading': el.get_text(strip=True), 'html_parts': []}
            elif current is not None and tag:
                current['html_parts'].append(str(el))

        if current is not None:
            sections.append(current)

        # Match sections to chapters
        updated = []
        skipped = []

        for sec in sections:
            heading = sec['heading']
            content_html = '\n'.join(sec['html_parts'])

            # Try to find matching chapter: heading starts with chapter_no
            matched_ch = None
            matched_no = None
            for ch_no, ch in chapter_map.items():
                if ch_no and heading.startswith(ch_no):
                    # Prefer longer (more specific) match
                    if matched_no is None or len(ch_no) > len(matched_no):
                        matched_ch = ch
                        matched_no = ch_no

            if matched_ch is None:
                skipped.append(heading)
                continue

            update_data = schemas.KMArticleUpdate(content=content_html)
            self.update_article(matched_ch.id, update_data)
            updated.append({'chapter_no': matched_no, 'title': matched_ch.title})

        return {
            'updated': updated,
            'skipped': skipped,
            'total_sections': len(sections),
        }

    ALLOWED_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp', '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.pptx', '.txt', '.csv', '.zip'}
    MAX_FILE_SIZE_MB = 10

    def upload_image(self, file: UploadFile) -> str:
        # Validate file extension
        ext = os.path.splitext(file.filename or '')[1].lower()
        if ext not in self.ALLOWED_EXTENSIONS:
            raise HTTPException(status_code=400, detail=f"File type '{ext}' is not allowed. Allowed: {', '.join(self.ALLOWED_EXTENSIONS)}")

        # Validate file size (read content first)
        contents = file.file.read()
        if len(contents) > self.MAX_FILE_SIZE_MB * 1024 * 1024:
            raise HTTPException(status_code=400, detail=f"File size exceeds {self.MAX_FILE_SIZE_MB}MB limit.")
        file.file.seek(0)  # Reset for writing

        # Resolve the base directory of the backend
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        upload_dir = os.path.join(base_dir, "uploads", "km")
        os.makedirs(upload_dir, exist_ok=True)

        file_uuid = str(uuid.uuid4())
        safe_filename = (file.filename or "file").replace(" ", "_")
        filename = f"km_{file_uuid}_{safe_filename}"
        file_path = os.path.join(upload_dir, filename)

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        return f"/api/files/download/km/{filename}"
